from docx import Document
from docx.shared import Cm, Pt

forfatter="forfatter"
bedriftsNavn="bedriftsnavn"
pris=8000
VEK=6500
mengde=500
FTK=300000


DB=pris-VEK
DBT=DB*mengde
overskudd=DBT-FTK
DG=round((DB*100)/pris, 2)
DPkroner=round(FTK/(DG/100))
DPenheter=round(FTK/DB)
STI=pris*mengde
SMkroner=round(STI-DPkroner)
SMenheter=round(mengde-DPenheter)
SMprosent=round(SMenheter*100/mengde, 2)
indent=1
#filNavn="DPAB.png"
#filNavn2="FAB.png"
#bildeWidth=15

#Input målsøk:
svarFelt=1
deltaT=0.032
faktor=DG
mål=45
variabel=VEK

faktorStr="DG"
variabelStr="VEK"

def målSøk(faktor, mål, variabel):
    global pris, VEK, mengde, FTK, DB, DBT, overskudd, DG, DPkroner, DPenheter, STI
    #Sjekker om det er rett faktor og variabel jeg regner på
    if faktor == DG and variabel == pris:
        pris3=pris
        #faktorStr="DG"
        #variabelStr="pris"
        #Sjekka om e allerede e i mål
        if DG == mål:
            return pris

        while True:
            DG2=DG
            if DG < mål:
                pris3 += deltaT
                #Endra på variabelen og sjekka på nytt
                DG = (pris3-VEK)*100/pris3
                #Sjekka om e har kommet meg noe nærmere, hvis ikke, return
                if mål - DG > mål - DG2:
                    return (pris3-deltaT)
            else:
                return round(pris3,1)
    
    #Sjekker om det er rett faktor og variabel jeg regner på
    if faktor == DG and variabel == VEK:
        VEK3=VEK
        if DG == mål:
            return VEK

        while True:
            DG2=DG
            #Sjekka om e allerede e i mål
            if DG < mål:
                VEK -= deltaT
                #Endra på variabelen og sjekka på nytt
                DG = (pris-VEK)*100/pris
                #Sjekka om e e nærmere, hvis ikke, return
                if mål - DG > mål - DG2:
                    return (VEK-deltaT)
            else:
                return round(VEK,1)
        

#Lager et dokument å jobbe i
document=Document()


#Lager en overskrift
document.add_heading("Dekningspunkt analyse av "+str(forfatter), 0)


#Enkelt paragraf:
p=document.add_paragraph("En dekningspunktsanalyse er en analyse over forskjellige aspekt innenfor en bedrift. Med en slik analyse kan vi lettere få et overblikk over kostnader, inntekter, salgsvolum, og hvor mye vi kan endre på disse verdiene og fremdeles få inn profitt.")


#Legg til tabell
data=[
    "Inndata:", "Navn:", str(bedriftsNavn),
    None, "Pris per enhet ekskl. mva:", str(pris),
    None, "Variable enhetskostnader ekskl. mva:", str(VEK),
    None, "Faste totale kostnader per år ekskl. mva:", str(FTK),
    None, "Produksjon/salg per år (enheter)", str(mengde),
    "Dekningsanalyse:", None, None,
    "Dekningsbidrag per enhet", None, str(DB),
    "Dekningsbidrag totalt ved "+str(mengde)+" enheter", None, str(DBT),
    "Dekningsgrad", None, str(DG)+"%",
    "Overskudd ved "+str(mengde)+" enheter", None, str(overskudd),
    "Dekningspunkt i kroner", None, str(DPkroner),
    "Dekningspunkt i enheter", None, str(DPenheter),
    "Sikkerhetsmargin i kroner ved "+str(mengde)+" enheter", None, str(SMkroner),
    "Sikkerhetsmargin i enheter ved "+str(mengde)+" enheter", None, str(SMenheter),
    "Sikkerhetsmargin i prosent (enheter) ved "+str(mengde)+" enheter", None, str(SMprosent)+"%"
]

table=document.add_table(rows=0, cols=3)
table.style="Light Grid Accent 5"

lenght=len(data)
for i in range(int(lenght/3)):
    x=i*3
    row=table.add_row().cells
    if data[x] != None:
        row[0].text=str(data[x])
    if data[x+1] != None:
        row[1].text=str(data[x+1])
    if data[x+2] != None:
        row[2].text=str(data[x+2])


#Forklaring
p1=document.add_paragraph("")
p1.add_run("Forklaringer:").bold=True
p1_format=p1.paragraph_format
p1_format.space_before=Pt(12)

#DB per enhet
p2=document.add_paragraph("")
p2.add_run("Dekningsbidrag per enhet:").bold=True
p2_format=p2.paragraph_format
p2_format.left_indent = Cm(indent)


p3tekst="Dekningsbidrag vil si hvor mye inntjeneste vi får når vi ser vekk ifra faste kostnader. Når vi skal ha det per enhet må vi dele på antall enheter. Her er dekningsbidraget per enhet "+ str(DB) +"kr for " + bedriftsNavn
p3=document.add_paragraph("")
p3.add_run(p3tekst).italic=True
p3_format=p3.paragraph_format
p3_format.left_indent = Cm(indent)

p4=document.add_paragraph("DB per enhet = Salgspris per enhet (P) - Variable enhetskostnader (VEK) = "+str(pris)+"kr - "+str(VEK)+"kr = "+str(DB)+"kr")
p4_format=p4.paragraph_format
p4_format.left_indent = Cm(indent)
#DB totalt
k2=document.add_paragraph("")
k2.add_run("Dekningsbidrag totalt:").bold=True
k2_format=k2.paragraph_format
k2_format.left_indent = Cm(indent)

k3tekst="Dekningsbidrag totalt er innetjenesten våres når vi ser vekk ifra faste kostnader."
k3=document.add_paragraph("")
k3.add_run(k3tekst).italic=True
k3_format=k3.paragraph_format
k3_format.left_indent = Cm(indent)

k4=document.add_paragraph("DBT = DB per enhet * mengde ="+str(DB)+"kr * "+str(mengde)+" enheter = "+str(DBT)+"kr")
k4_format=k4.paragraph_format
k4_format.left_indent = Cm(indent)
#Dekningsgrad
i2=document.add_paragraph("")
i2.add_run("Dekningsgrad:").bold=True
i2_format=i2.paragraph_format
i2_format.left_indent = Cm(indent)

i3tekst="Dekningsgrad er hvor mye av salgsprisen som blir dekningsbidrag."
i3=document.add_paragraph("")
i3.add_run(i3tekst).italic=True
i3_format=i3.paragraph_format
i3_format.left_indent = Cm(indent)

i4=document.add_paragraph("DG = (DP per enhet *100%)/pris = ("+str(DB)+" * 100%)/"+str(pris)+" = "+str(DG)+"%")
i4_format=i4.paragraph_format
i4_format.left_indent = Cm(indent)
#Overskudd
o2=document.add_paragraph("")
o2.add_run("Overskudd:").bold=True
o2_format=o2.paragraph_format
o2_format.left_indent = Cm(indent)

o3tekst="Overskudd er hvor mye penger bedriften vil sitte igjen med etter salg og kostnader. Disse pengene kan "+bedriftsNavn+" bruke til forskjellige ting, inkludert å sette inn igjen i bedriften, og utbytte til investorer."
o3=document.add_paragraph("")
o3.add_run(o3tekst).italic=True
o3_format=o3.paragraph_format
o3_format.left_indent = Cm(indent)

o4=document.add_paragraph("Overskudd = DBT - Faste totale kostnader (FTK) = "+str(DBT)+"kr - "+str(FTK)+"kr = "+str(overskudd)+"kr")
o4_format=o4.paragraph_format
o4_format.left_indent = Cm(indent)
#Dekningspunkt i kroner
l2=document.add_paragraph("")
l2.add_run("Dekningspunkt i kroner:").bold=True
l2_format=l2.paragraph_format
l2_format.left_indent = Cm(indent)

l3tekst="Dekningspunkt i kroner er hvor mye innetjeneste vi må ha av salget for å tjene inn igjen de faste kostnadene."
l3=document.add_paragraph("")
l3.add_run(l3tekst).italic=True
l3_format=l3.paragraph_format
l3_format.left_indent = Cm(indent)

l4=document.add_paragraph("DP i kroner = FTK/DG = "+str(FTK)+"kr / "+str(DG)+"% = "+str(DPkroner)+"kr")
l4_format=l4.paragraph_format
l4_format.left_indent = Cm(indent)
#Dekningspunkt i enheter
u2=document.add_paragraph("")
u2.add_run("Dekningspunkt i enheter:").bold=True
u2_format=u2.paragraph_format
u2_format.left_indent = Cm(indent)

u3tekst="Dekningspunkt i enheter er hvor mange enheter vi må selge for å dekke inn de faste kostnadene."
u3=document.add_paragraph("")
u3.add_run(u3tekst).italic=True
u3_format=u3.paragraph_format
u3_format.left_indent = Cm(indent)

u4=document.add_paragraph("DP i enheter = FTK / DB per enhet = "+str(FTK)+"kr / "+str(DB)+"kr = "+str(DPenheter))
u4_format=u4.paragraph_format
u4_format.left_indent = Cm(indent)
#Sum total inntekt
j2=document.add_paragraph("")
j2.add_run("Sum total inntekt:").bold=True
j2_format=j2.paragraph_format
j2_format.left_indent = Cm(indent)

j3tekst="Sum total inntekt er inntekten vi får av salg, uten å ta i bekostning de faste kostnadene."
j3=document.add_paragraph("")
j3.add_run(j3tekst).italic=True
j3_format=j3.paragraph_format
j3_format.left_indent = Cm(indent)

j4=document.add_paragraph("Sum total inntekter = pris * mengde = "+str(pris)+"kr * "+str(mengde)+" enheter = "+str(STI)+"kr")
j4_format=j4.paragraph_format
j4_format.left_indent = Cm(indent)
#Sikkerhetsmargin kroner
h2=document.add_paragraph("")
h2.add_run("Sikkerhetsmargin i kroner:").bold=True
h2_format=h2.paragraph_format
h2_format.left_indent = Cm(indent)

h3tekst="Sikkerhetsmargin i kroner er hvor mye penger vi kan tape og fremdeles gå i null."
h3=document.add_paragraph("")
h3.add_run(h3tekst).italic=True
h3_format=h3.paragraph_format
h3_format.left_indent = Cm(indent)

h4=document.add_paragraph("SM i kroner = STI - DP i kroner = "+str(STI)+"kr - "+str(DPkroner)+"kr = "+str(SMkroner)+"kr")
h4_format=h4.paragraph_format
h4_format.left_indent = Cm(indent)
#SM enheter
y2=document.add_paragraph("")
y2.add_run("Sikkerhetsmargin i enheter:").bold=True
y2_format=y2.paragraph_format
y2_format.left_indent = Cm(indent)

y3tekst="Sikkerhetsmargin i enheter er hvor mange færre enheter vi kan selge, og fremdeles gå i null."
y3=document.add_paragraph("")
y3.add_run(y3tekst).italic=True
y3_format=y3.paragraph_format
y3_format.left_indent = Cm(indent)

y4=document.add_paragraph("SM i enheter = Mengde - Dp i enheter = "+str(mengde)+" enheter - "+str(DPenheter)+" enheter = "+str(SMenheter)+" enheter")
y4_format=y4.paragraph_format
y4_format.left_indent = Cm(indent)
#SM prosent (enheter)
g2=document.add_paragraph("")
g2.add_run("Sikkerhetsmargin i prosent (enheter):").bold=True
g2_format=g2.paragraph_format
g2_format.left_indent = Cm(indent)

g3tekst="Sikkerhetsmargin i prosent er hvor mange færre enheter "+bedriftsNavn+" kan selge og fremdeles gå i null."
g3=document.add_paragraph("")
g3.add_run(g3tekst).italic=True
g3_format=g3.paragraph_format
g3_format.left_indent = Cm(indent)

g4=document.add_paragraph("SM i prosent = (SM i enheter * 100%) / Mengde = ("+str(SMenheter)+" enheter * 100%) / "+str(mengde)+" enheter = "+str(SMprosent)+"%")
g4_format=g4.paragraph_format
g4_format.left_indent = Cm(indent)
#Følsomhetsanalyse, kan prøve deg på en if større/mindre enn x. e ville spurt anna lol
#Bibliografi? idk man

f1tekst="En følsomhetsanalyse er en analyse som sier oss hvor mye vi kan endre på de ulike tallene, og fremdeles være en lønnsom bedrift. Lav prosent kan være en høyere risiko, og motsatt."
f1=document.add_paragraph(f1tekst)
f1_format=f1.paragraph_format
f1_format.space_before=Pt(42)


#Variabler følsomhetsanalyse:
#Pris, VEK, mengde, FTK, DB, DBT, overskudd, DG, DPkroner, DPenheter, STI, SMkroner, SMenheter, SMprosent
pris2=(FTK/mengde)+VEK
marginPris=-round((1-(pris2/pris))*100,1)

VEK2=pris-(FTK/mengde)
marginVEK=-round((1-(VEK2/VEK))*100,1)

FTK2=mengde*(pris-VEK)
marginFTK=-round((1-(FTK2/FTK))*100,1)

mengde2=FTK/(pris-VEK)
marginMengde=-round((1-(mengde2/mengde))*100,1)


#Tabell følsomhetsanalyse
kritiskPris=round(pris2-pris, 2)
kritiskVEK=round((VEK2-VEK),2)
kritiskFTK=round((FTK2-FTK),2)
kritiskMengde=round((mengde2-mengde),2)

dataF=[
    "Følsomhetsanalyse", None, None, None, None,
    "Variabel", "Verdi", "Kritisk verdi", "Margin", "Margin %",
    "Pris", str(pris), str(round(pris2,2)), str(kritiskPris), str(marginPris)+"%",
    "Variable kostnader", str(VEK), str(round(VEK2,2)), str(kritiskVEK), str(marginVEK)+"%",
    "Faste kostnader", str(FTK), str(round(FTK2,2)), str(kritiskFTK), str(marginFTK)+"%",
    "Mengde", str(mengde), str(round(mengde2,2)), str(kritiskMengde), str(marginMengde)+"%"
]

tableF=document.add_table(rows=0, cols=5)
tableF.style="Light Grid Accent 5"

lenght=len(dataF)
for i in range(int(lenght/5)):
    x=i*5
    row=tableF.add_row().cells
    if dataF[x] != None:
        row[0].text=str(dataF[x])
    if dataF[x+1] != None:
        row[1].text=str(dataF[x+1])
    if dataF[x+2] != None:
        row[2].text=str(dataF[x+2])
    if dataF[x+3] != None:
        row[3].text=str(dataF[x+3])
    if dataF[x+4] != None:
        row[4].text=str(dataF[x+4])




#Margin pris
if abs(marginPris) <=5:
    tekstPris=str(marginPris)+"%, dette er en relativt lav margin, med relativt høy risiko. Her er det ikke anbefalt for "+bedriftsNavn+" å endre på prisen for å konkurrere med andre bedrifter. "
if 5 < abs(marginPris) < 15:
    tekstPris=str(marginPris)+"%, dette er en moderat margin, uten en spesielt stor risiko. Her er det litt rom for "+bedriftsNavn+" å endre prisen, men de må være forsiktig. "
if 15 <= abs(marginPris):
    tekstPris=str(marginPris)+"%, dette er en relativt høy margin, som betyr liten risiko. Her kan "+bedriftsNavn+" endre en del på prisen uten å være redd for å gå i tap. "
#Margin VEK
if abs(marginVEK) <=5:
    tekstVEK=str(marginVEK)+"%, en såpass lav VEK margin kan være problematisk om de variable kostnadene stiger uforventet, og er noe "+bedriftsNavn+" burde være overvåken over. "
if 5 < abs(marginVEK) < 15:
    tekstVEK=str(marginVEK)+"%, en moderat margin som dette vil ikke være problematisk for "+bedriftsNavn+" med mindre det kommer større endringer i variable kostnader. "
if 15 <= abs(marginVEK):
    tekstVEK=str(marginVEK)+"%, en såpass høy margin har veldig liten risiko og "+bedriftsNavn+" trenger ikke være bekymret for en endring av variable kostnader med det første. "
#Margin FTK
if abs(marginFTK) <=10:
    tekstFTK=str(marginFTK)+"%, ved faste kostnader er dette en relativt lav prosent og "+bedriftsNavn+" må passe seg for eventuelle endringer i de fastekostnadene. "
if 10 < abs(marginFTK) < 30:
    tekstFTK=str(marginFTK)+"%, noe som er en relativt normal margin for faste kostnader. Her trenger ikke "+bedriftsNavn+" å bekymre seg i så stor grad for eventuelle endringer. "
if 30 <= abs(marginFTK):
    tekstFTK=str(marginFTK)+"%, noe som er en relativt høy margin for faste kostnader. Her kan "+bedriftsNavn+" føle seg trygg på at det skal drastiske endringer i FTK før de går i tap. "
#Margin mengde
if abs(marginMengde) <= 10:
    tekstMengde=str(marginMengde)+"%, dette vil si at "+bedriftsNavn+" har en relativt lav margin på enheter solgt og burde være fosiktig med handlinger som kan redusere salget. "
if 10 < abs(marginMengde) < 25:
    tekstMengde=str(marginMengde)+"%, dette vil si at "+bedriftsNavn+" har en moderat margin på enheter solgt. "
if 25 <= abs(marginMengde):
    tekstMengde=str(marginMengde)+"%, dette vil si at "+bedriftsNavn+" har en relativt stor margin på enheter solgt, og trenger ikke bekymre seg for å gå i null selv om enheter solgt går ned. "

f2tekst="Her kan vi se at prisen har en margin på "+tekstPris+"Ut ifra tabellen ser vi at de variable enhetskostnadene har en margin på "+tekstVEK+" Vi finner også marginen til de faste kostnadene som er på "+tekstFTK+" Videre kan vi se at mengden har en margin på "+tekstMengde
f2=document.add_paragraph(f2tekst)
f2_format=f2.paragraph_format
f2_format.space_before=Pt(12)


if svarFelt == 1:
    målSvar=målSøk(faktor, mål, variabel)
    målSøkTekst2="Når man ønsker å endre på "+str(faktorStr)+" kan det være flere variabler i spill, her velger jeg å endre på "+str(variabelStr)+" og vi får da at "
    målSøkTekst=målSøkTekst2+"hvis "+bedriftsNavn+" ønsker en "+str(faktorStr)+ " på "+str(mål)+" må de endre "+str(variabelStr)+" til "+str(målSvar)
    målS=document.add_paragraph(str(målSøkTekst))
    målS_format=målS.paragraph_format
    målS_format.space_before=Pt(28)


biblo=document.add_paragraph("Bibliografi")
biblo_format=biblo.paragraph_format
biblo_format.space_before=Pt(52)


kildeTekst="Bringsli, Øystein (2021, 12 3). Dekningspunkt og følsomhetsanalyse i Python."
kildeT=document.add_paragraph("")
kildeT.add_run(kildeTekst).italic=True


#Lagre dokumentet:
document.save("Dekningspunksanalyse med tabell "+bedriftsNavn+".docx")
